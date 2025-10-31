#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Markdown转政府公文格式Word文档 - 稳定生产版本
功能：一键将Markdown文件转换为符合政府公文格式标准的Word文档
版本：1.0.0
"""

import docx
from docx.shared import Pt, Mm, RGBColor
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
import sys
import re
from pathlib import Path

# ==================== 常量定义 ====================
# 字体定义
FONT_FANGSONG_GB2312 = '仿宋_GB2312'
FONT_KAITI_GB2312 = '楷体_GB2312'
FONT_HEITI = '黑体'
FONT_XIAOBIAOSONG = '方正小标宋简体'

# 字号定义（按国标）
SIZE_ERHAO = Pt(22)      # 二号：22磅
SIZE_SANHAO = Pt(16)     # 三号：16磅
SIZE_XIAOSAN = Pt(15)    # 小三：15磅

# 行距和缩进
LINE_SPACING_28_8 = Pt(28.8)   # 固定行距28.8磅
FIRST_LINE_INDENT_32 = Pt(32)  # 首行缩进2字符（三号字体约32磅）

# 页边距（单位：毫米）
MARGIN_TOP = Mm(37)
MARGIN_BOTTOM = Mm(35)
MARGIN_LEFT = Mm(28)
MARGIN_RIGHT = Mm(26)


# ==================== Markdown模式识别 ====================
# 标题识别（支持多种Markdown标记）
MD_H1_PATTERN = re.compile(r'^#\s+(.+)$')           # # 主标题
MD_H2_PATTERN = re.compile(r'^##\s+(.+)$')          # ## 一级标题
MD_H3_PATTERN = re.compile(r'^###\s+(.+)$')         # ### 二级标题
MD_H4_PATTERN = re.compile(r'^####\s+(.+)$')        # #### 三级标题

# 序号格式处理正则（用于移除序号后的空格）
NUMBER_SPACE_PATTERN = re.compile(r'^([一二三四五六七八九十]+、|\d+\.|（[一二三四五六七八九十]+）)\s+')

# 列表项识别
MD_LIST_ITEM_PATTERN = re.compile(r'^\s*[-*+☑]\s+(.+)$')

# 行内格式识别
MD_BOLD_PATTERN = re.compile(r'\*\*(.+?)\*\*')      # **加粗**
MD_ITALIC_PATTERN = re.compile(r'\*(.+?)\*')        # *斜体*
MD_CODE_PATTERN = re.compile(r'`(.+?)`')            # `代码`

# 分隔线
MD_SEPARATOR_PATTERN = re.compile(r'^[-*_]{3,}$')

# 表格识别
MD_TABLE_SEPARATOR = re.compile(r'^\s*\|?\s*[-:]+\s*\|')


# ==================== 工具函数 ====================
def set_run_format(run, font_name, font_size, bold=False, italic=False, color=None):
    """
    统一设置Run的格式
    
    参数:
        run: docx.text.run.Run对象
        font_name: 字体名称
        font_size: 字号（Pt对象）
        bold: 是否加粗
        italic: 是否斜体
        color: 颜色（RGBColor对象，可选）
    """
    run.font.name = font_name
    run.font.element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = font_size
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color


def apply_paragraph_format(p_format, 
                          alignment=WD_ALIGN_PARAGRAPH.JUSTIFY_LOW,
                          indent=FIRST_LINE_INDENT_32, 
                          line_spacing=LINE_SPACING_28_8,
                          space_before=Pt(0),
                          space_after=Pt(0)):
    """
    统一应用段落格式
    
    参数:
        p_format: 段落格式对象
        alignment: 对齐方式
        indent: 首行缩进
        line_spacing: 行距
        space_before: 段前间距
        space_after: 段后间距
    """
    p_format.alignment = alignment
    p_format.first_line_indent = indent
    p_format.line_spacing = line_spacing
    p_format.space_before = space_before
    p_format.space_after = space_after


def parse_inline_format(text):
    """
    解析文本中的行内格式（加粗、斜体等）
    
    返回: [(文本片段, 是否加粗, 是否斜体), ...]
    """
    segments = []
    pos = 0
    
    # 匹配所有行内格式标记
    patterns = [
        (MD_BOLD_PATTERN, True, False),   # 加粗
        (MD_ITALIC_PATTERN, False, True),  # 斜体
    ]
    
    while pos < len(text):
        matched = False
        
        # 尝试匹配加粗（优先级更高）
        bold_match = MD_BOLD_PATTERN.match(text, pos)
        if bold_match:
            segments.append((bold_match.group(1), True, False))
            pos = bold_match.end()
            matched = True
            continue
        
        # 尝试匹配斜体
        italic_match = MD_ITALIC_PATTERN.match(text, pos)
        if italic_match:
            segments.append((italic_match.group(1), False, True))
            pos = italic_match.end()
            matched = True
            continue
        
        # 没有匹配到格式标记，添加普通字符
        if not matched:
            # 找到下一个可能的格式标记位置
            next_pos = len(text)
            for pattern, _, _ in patterns:
                match = pattern.search(text, pos)
                if match and match.start() < next_pos:
                    next_pos = match.start()
            
            # 添加普通文本片段
            if next_pos > pos:
                segments.append((text[pos:next_pos], False, False))
                pos = next_pos
            else:
                segments.append((text[pos], False, False))
                pos += 1
    
    return segments if segments else [(text, False, False)]


def remove_number_space(text):
    """
    移除序号后的空格
    例如："一、 总体要求" -> "一、总体要求"
          "1. 字体要求" -> "1.字体要求"
          "（一） 加强培训" -> "（一）加强培训"
    """
    return NUMBER_SPACE_PATTERN.sub(r'\1', text)


def clean_markdown_marks(text):
    """
    清理文本中的Markdown格式标记（用于标题等不需要格式的地方）
    移除：**加粗**、*斜体*、`代码` 等标记
    
    例如："**一级标题**" -> "一级标题"
          "*重要*内容" -> "重要内容"
    """
    # 移除加粗标记 **text**
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    # 移除斜体标记 *text*
    text = re.sub(r'\*(.+?)\*', r'\1', text)
    # 移除代码标记 `text`
    text = re.sub(r'`(.+?)`', r'\1', text)
    # 移除删除线 ~~text~~
    text = re.sub(r'~~(.+?)~~', r'\1', text)
    return text


def add_formatted_text(paragraph, text, base_font, base_size):
    """
    向段落添加带格式的文本（处理加粗、斜体等）
    
    参数:
        paragraph: 段落对象
        text: 文本内容
        base_font: 基础字体
        base_size: 基础字号
    """
    # 移除序号后的空格
    text = remove_number_space(text)
    
    segments = parse_inline_format(text)
    for content, bold, italic in segments:
        if content:
            run = paragraph.add_run(content)
            set_run_format(run, base_font, base_size, bold=bold, italic=italic)


def setup_page_margins(doc):
    """设置页面边距"""
    for section in doc.sections:
        section.top_margin = MARGIN_TOP
        section.bottom_margin = MARGIN_BOTTOM
        section.left_margin = MARGIN_LEFT
        section.right_margin = MARGIN_RIGHT


def parse_table_row(line):
    """
    解析Markdown表格行
    返回单元格列表
    """
    # 移除首尾的 |
    line = line.strip()
    if line.startswith('|'):
        line = line[1:]
    if line.endswith('|'):
        line = line[:-1]
    
    # 分割单元格
    cells = [cell.strip() for cell in line.split('|')]
    return cells


def add_table_to_doc(doc, table_data):
    """
    向Word文档添加表格
    
    参数:
        doc: Word文档对象
        table_data: 表格数据 [[header1, header2, ...], [row1col1, row1col2, ...], ...]
    """
    if not table_data or len(table_data) < 2:
        return
    
    # 创建表格
    rows = len(table_data)
    cols = len(table_data[0])
    table = doc.add_table(rows=rows, cols=cols)
    
    # 设置表格样式
    table.style = 'Table Grid'
    
    # 填充表格内容
    for i, row_data in enumerate(table_data):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            if j < len(row.cells):
                cell = row.cells[j]
                # 设置单元格文本
                cell.text = cell_text
                
                # 设置单元格格式
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    paragraph.paragraph_format.space_before = Pt(3)
                    paragraph.paragraph_format.space_after = Pt(3)
                    
                    for run in paragraph.runs:
                        # 表头加粗
                        if i == 0:
                            set_run_format(run, FONT_HEITI, SIZE_SANHAO, bold=True)
                        else:
                            set_run_format(run, FONT_FANGSONG_GB2312, SIZE_SANHAO)
    
    # 表格后添加空行
    doc.add_paragraph()


# ==================== 核心转换函数 ====================
def convert_markdown_to_gov_docx(md_path, docx_path):
    """
    将Markdown文件转换为政府公文格式的Word文档
    
    参数:
        md_path: Markdown文件路径
        docx_path: 输出的Word文档路径
    
    返回:
        成功返回True，失败返回False
    """
    try:
        # 检查输入文件是否存在
        md_file = Path(md_path)
        if not md_file.exists():
            print(f"❌ 错误：输入文件不存在: {md_path}")
            return False
        
        # 读取Markdown内容
        with open(md_path, 'r', encoding='utf-8') as f:
            lines = f.readlines()
        
        # 创建Word文档
        doc = docx.Document()
        
        # 设置页面边距
        setup_page_margins(doc)
        
        # 标记第一个标题（作为主标题）
        is_first_heading = True
        
        # 表格处理状态
        in_table = False
        table_data = []
        
        # 逐行处理
        i = 0
        while i < len(lines):
            line = lines[i]
            text = line.strip()
            
            # 跳过空行和分隔线
            if not text or MD_SEPARATOR_PATTERN.match(text):
                i += 1
                continue
            
            # ============ 检测表格 ============
            if '|' in text and i + 1 < len(lines):
                next_line = lines[i + 1].strip()
                # 检查下一行是否是表格分隔符
                if MD_TABLE_SEPARATOR.match(next_line):
                    # 开始解析表格
                    in_table = True
                    table_data = []
                    
                    # 表头
                    table_data.append(parse_table_row(text))
                    i += 1  # 跳过分隔符行
                    i += 1
                    
                    # 读取表格数据行
                    while i < len(lines):
                        row_text = lines[i].strip()
                        if not row_text or '|' not in row_text:
                            break
                        table_data.append(parse_table_row(row_text))
                        i += 1
                    
                    # 添加表格到文档
                    add_table_to_doc(doc, table_data)
                    in_table = False
                    continue
            
            # 创建新段落
            para = doc.add_paragraph()
            
            # ============ 1. 主标题（第一个一级标题）============
            match = MD_H1_PATTERN.match(text)
            if match and is_first_heading:
                is_first_heading = False
                title_text = clean_markdown_marks(match.group(1))  # 清理格式标记
                apply_paragraph_format(
                    para.paragraph_format,
                    alignment=WD_ALIGN_PARAGRAPH.CENTER,
                    indent=Pt(0),
                    line_spacing=LINE_SPACING_28_8
                )
                run = para.add_run(title_text)
                set_run_format(run, FONT_XIAOBIAOSONG, SIZE_ERHAO, bold=True)
                
                # 主标题后添加空行
                doc.add_paragraph()
                i += 1
                continue
            
            # ============ 2. 一级标题（## 开头）============
            match = MD_H2_PATTERN.match(text)
            if match:
                heading_text = clean_markdown_marks(match.group(1))  # 清理格式标记
                apply_paragraph_format(para.paragraph_format)
                run = para.add_run(heading_text)
                set_run_format(run, FONT_HEITI, SIZE_SANHAO)
                i += 1
                continue
            
            # ============ 3. 二级标题（### 开头）============
            match = MD_H3_PATTERN.match(text)
            if match:
                heading_text = clean_markdown_marks(match.group(1))  # 清理格式标记
                apply_paragraph_format(para.paragraph_format)
                run = para.add_run(heading_text)
                set_run_format(run, FONT_KAITI_GB2312, SIZE_SANHAO, bold=True)
                i += 1
                continue
            
            # ============ 4. 三级标题（#### 开头）============
            match = MD_H4_PATTERN.match(text)
            if match:
                heading_text = clean_markdown_marks(match.group(1))  # 清理格式标记
                apply_paragraph_format(para.paragraph_format)
                run = para.add_run(heading_text)
                set_run_format(run, FONT_KAITI_GB2312, SIZE_SANHAO)
                i += 1
                continue
            
            # ============ 5. 列表项 ============
            match = MD_LIST_ITEM_PATTERN.match(text)
            if match:
                list_text = match.group(1)
                apply_paragraph_format(para.paragraph_format)
                add_formatted_text(para, list_text, FONT_FANGSONG_GB2312, SIZE_SANHAO)
                i += 1
                continue
            
            # ============ 6. 普通正文 ============
            apply_paragraph_format(para.paragraph_format)
            add_formatted_text(para, text, FONT_FANGSONG_GB2312, SIZE_SANHAO)
            
            i += 1
        
        # 保存文档
        output_path = Path(docx_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(docx_path)
        
        print(f"✅ 转换成功！")
        print(f"   输入: {md_path}")
        print(f"   输出: {docx_path}")
        return True
        
    except Exception as e:
        print(f"❌ 转换失败: {e}")
        import traceback
        traceback.print_exc()
        return False


# ==================== 命令行入口 ====================
def main():
    """命令行主函数"""
    if len(sys.argv) != 3:
        print("\n" + "="*60)
        print("Markdown转政府公文格式Word文档工具 v1.0.0")
        print("="*60)
        print("\n使用方法:")
        print(f"  python {Path(__file__).name} <输入.md文件> <输出.docx文件>")
        print("\n示例:")
        print(f"  python {Path(__file__).name} report.md report_formatted.docx")
        print("\n支持的Markdown语法:")
        print("  # 主标题        -> 方正小标宋简体 22磅 加粗 居中")
        print("  ## 一级标题     -> 黑体 16磅")
        print("  ### 二级标题    -> 楷体_GB2312 16磅 加粗")
        print("  #### 三级标题   -> 楷体_GB2312 16磅")
        print("  - 列表项        -> 仿宋_GB2312 16磅")
        print("  正文            -> 仿宋_GB2312 16磅")
        print("  **加粗**        -> 加粗")
        print("  *斜体*          -> 斜体")
        print("\n注意事项:")
        print("  - 请确保系统已安装所需字体（仿宋_GB2312、楷体_GB2312等）")
        print("  - macOS用户可在'字体册'中检查字体")
        print("="*60 + "\n")
        sys.exit(1)
    
    input_file = sys.argv[1]
    output_file = sys.argv[2]
    
    success = convert_markdown_to_gov_docx(input_file, output_file)
    sys.exit(0 if success else 1)


if __name__ == '__main__':
    main()
